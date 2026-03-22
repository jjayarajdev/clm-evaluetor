"""
Generate test contract documents for E2E testing.
Creates realistic contracts with various clauses, obligations, and risk factors.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def create_msa():
    """Create a Master Service Agreement with various clauses."""
    doc = Document()

    # Title
    title = doc.add_heading("MASTER SERVICE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parties
    doc.add_paragraph()
    add_paragraph(doc, "This Master Service Agreement (the 'Agreement') is entered into as of January 15, 2025 (the 'Effective Date') by and between:")
    doc.add_paragraph()
    add_paragraph(doc, "Acme Corporation, a Delaware corporation with its principal place of business at 123 Business Park, San Francisco, CA 94102 ('Client')")
    doc.add_paragraph()
    add_paragraph(doc, "AND")
    doc.add_paragraph()
    add_paragraph(doc, "TechVendor Solutions Inc., a California corporation with its principal place of business at 456 Innovation Drive, Palo Alto, CA 94301 ('Vendor')")

    # Recitals
    add_heading(doc, "RECITALS", 1)
    add_paragraph(doc, "WHEREAS, Vendor is engaged in the business of providing software development, IT consulting, and managed services;")
    add_paragraph(doc, "WHEREAS, Client desires to engage Vendor to provide certain services as described herein;")
    add_paragraph(doc, "NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, the parties agree as follows:")

    # Article 1 - Services
    add_heading(doc, "ARTICLE 1: SERVICES", 1)
    add_paragraph(doc, "1.1 Scope of Services. Vendor shall provide the services described in each Statement of Work ('SOW') executed by both parties. Each SOW shall be incorporated into and governed by this Agreement.")
    add_paragraph(doc, "1.2 Service Levels. Vendor shall maintain a minimum system uptime of 99.5% measured monthly. Vendor shall respond to critical issues within 2 hours and resolve them within 8 hours.")
    add_paragraph(doc, "1.3 Performance Standards. Vendor shall perform all services in a professional and workmanlike manner consistent with industry standards.")

    # Article 2 - Term and Termination
    add_heading(doc, "ARTICLE 2: TERM AND TERMINATION", 1)
    add_paragraph(doc, "2.1 Term. This Agreement shall commence on the Effective Date and continue for a period of three (3) years, unless earlier terminated in accordance with this Article 2. This Agreement shall automatically renew for successive one (1) year periods unless either party provides written notice of non-renewal at least ninety (90) days prior to the end of the then-current term.")
    add_paragraph(doc, "2.2 Termination for Convenience. Either party may terminate this Agreement for any reason upon sixty (60) days prior written notice to the other party.")
    add_paragraph(doc, "2.3 Termination for Cause. Either party may terminate this Agreement immediately upon written notice if the other party: (a) materially breaches this Agreement and fails to cure such breach within thirty (30) days of receiving written notice; or (b) becomes insolvent or files for bankruptcy.")

    # Article 3 - Payment
    add_heading(doc, "ARTICLE 3: PAYMENT TERMS", 1)
    add_paragraph(doc, "3.1 Fees. Client shall pay Vendor the fees set forth in each SOW. The base monthly fee is $125,000.00 USD.")
    add_paragraph(doc, "3.2 Payment Terms. All invoices are due and payable within thirty (30) days of receipt. Late payments shall accrue interest at a rate of 1.5% per month.")
    add_paragraph(doc, "3.3 Expenses. Client shall reimburse Vendor for all pre-approved, reasonable travel and out-of-pocket expenses incurred in connection with the services.")

    # Article 4 - Confidentiality
    add_heading(doc, "ARTICLE 4: CONFIDENTIALITY", 1)
    add_paragraph(doc, "4.1 Definition. 'Confidential Information' means any non-public information disclosed by either party to the other, whether orally or in writing, that is designated as confidential or that reasonably should be understood to be confidential.")
    add_paragraph(doc, "4.2 Obligations. Each party agrees to: (a) maintain the confidentiality of the other party's Confidential Information; (b) not disclose such information to third parties without prior written consent; and (c) use such information only for purposes of this Agreement.")
    add_paragraph(doc, "4.3 Duration. The confidentiality obligations shall survive termination of this Agreement for a period of five (5) years.")

    # Article 5 - Intellectual Property
    add_heading(doc, "ARTICLE 5: INTELLECTUAL PROPERTY", 1)
    add_paragraph(doc, "5.1 Client Materials. Client shall retain all rights to any materials, data, or intellectual property provided to Vendor.")
    add_paragraph(doc, "5.2 Work Product. All work product created by Vendor specifically for Client under this Agreement shall be owned by Client upon full payment.")
    add_paragraph(doc, "5.3 Vendor Tools. Vendor retains all rights to its pre-existing tools, methodologies, and know-how.")

    # Article 6 - Liability
    add_heading(doc, "ARTICLE 6: LIABILITY AND INDEMNIFICATION", 1)
    add_paragraph(doc, "6.1 Limitation of Liability. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES. VENDOR'S TOTAL LIABILITY SHALL NOT EXCEED THE AMOUNTS PAID BY CLIENT IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM.")
    add_paragraph(doc, "6.2 Indemnification. Vendor shall indemnify, defend, and hold harmless Client from any third-party claims arising from Vendor's negligence, willful misconduct, or breach of this Agreement.")
    add_paragraph(doc, "6.3 Insurance. Vendor shall maintain commercial general liability insurance with coverage of at least $2,000,000 per occurrence and professional liability insurance of at least $5,000,000.")

    # Article 7 - Data Protection
    add_heading(doc, "ARTICLE 7: DATA PROTECTION", 1)
    add_paragraph(doc, "7.1 Compliance. Vendor shall comply with all applicable data protection laws, including GDPR and CCPA.")
    add_paragraph(doc, "7.2 Security. Vendor shall implement appropriate technical and organizational measures to protect Client data, including encryption at rest and in transit.")
    add_paragraph(doc, "7.3 Breach Notification. Vendor shall notify Client of any data breach within 24 hours of discovery.")

    # Article 8 - General
    add_heading(doc, "ARTICLE 8: GENERAL PROVISIONS", 1)
    add_paragraph(doc, "8.1 Governing Law. This Agreement shall be governed by the laws of the State of California.")
    add_paragraph(doc, "8.2 Dispute Resolution. Any disputes shall be resolved through binding arbitration in San Francisco, California.")
    add_paragraph(doc, "8.3 Entire Agreement. This Agreement constitutes the entire agreement between the parties and supersedes all prior negotiations and agreements.")
    add_paragraph(doc, "8.4 Amendment. This Agreement may only be amended by a written instrument signed by both parties.")
    add_paragraph(doc, "8.5 Assignment. Neither party may assign this Agreement without the prior written consent of the other party.")

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature blocks
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "ACME CORPORATION"
    table.cell(0, 1).text = "TECHVENDOR SOLUTIONS INC."
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name: John Smith"
    table.cell(2, 1).text = "Name: Jane Doe"
    table.cell(3, 0).text = "Title: Chief Executive Officer"
    table.cell(3, 1).text = "Title: President"

    # Save
    filepath = os.path.join(OUTPUT_DIR, "Acme_TechVendor_MSA_2025.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_nda():
    """Create a Non-Disclosure Agreement."""
    doc = Document()

    # Title
    title = doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_paragraph(doc, "This Mutual Non-Disclosure Agreement ('Agreement') is entered into as of February 1, 2025 ('Effective Date') by and between:")
    doc.add_paragraph()
    add_paragraph(doc, "Acme Corporation ('Acme'), and")
    add_paragraph(doc, "GlobalPartners Ltd., a UK company ('Partner')")
    doc.add_paragraph()
    add_paragraph(doc, "collectively referred to as the 'Parties' and individually as a 'Party'.")

    # Purpose
    add_heading(doc, "1. PURPOSE", 1)
    add_paragraph(doc, "The Parties wish to explore a potential business relationship concerning joint product development and market expansion ('Purpose'). In connection with this Purpose, each Party may disclose certain confidential and proprietary information to the other Party.")

    # Definition
    add_heading(doc, "2. DEFINITION OF CONFIDENTIAL INFORMATION", 1)
    add_paragraph(doc, "'Confidential Information' means any and all non-public information, in any form, disclosed by one Party ('Disclosing Party') to the other Party ('Receiving Party'), including but not limited to: trade secrets, business plans, customer lists, financial data, technical data, product designs, source code, algorithms, and marketing strategies.")

    # Obligations
    add_heading(doc, "3. OBLIGATIONS", 1)
    add_paragraph(doc, "3.1 The Receiving Party shall: (a) hold all Confidential Information in strict confidence; (b) not disclose Confidential Information to any third party without prior written consent; (c) use Confidential Information solely for the Purpose; (d) limit access to Confidential Information to employees with a need to know.")
    add_paragraph(doc, "3.2 The Receiving Party shall protect Confidential Information using at least the same degree of care it uses to protect its own confidential information, but in no event less than reasonable care.")

    # Exclusions
    add_heading(doc, "4. EXCLUSIONS", 1)
    add_paragraph(doc, "Confidential Information does not include information that: (a) is or becomes publicly available through no fault of the Receiving Party; (b) was known to the Receiving Party prior to disclosure; (c) is independently developed by the Receiving Party; (d) is rightfully received from a third party without restriction.")

    # Term
    add_heading(doc, "5. TERM", 1)
    add_paragraph(doc, "This Agreement shall remain in effect for two (2) years from the Effective Date. The confidentiality obligations shall survive termination for a period of three (3) years.")

    # Return of Materials
    add_heading(doc, "6. RETURN OF MATERIALS", 1)
    add_paragraph(doc, "Upon termination of this Agreement or upon request by the Disclosing Party, the Receiving Party shall promptly return or destroy all Confidential Information and certify such destruction in writing within ten (10) days.")

    # No License
    add_heading(doc, "7. NO LICENSE", 1)
    add_paragraph(doc, "Nothing in this Agreement grants any rights or licenses to any patents, copyrights, trademarks, or other intellectual property rights of either Party.")

    # Remedies
    add_heading(doc, "8. REMEDIES", 1)
    add_paragraph(doc, "Each Party acknowledges that a breach of this Agreement may cause irreparable harm for which monetary damages would be inadequate. Therefore, the non-breaching Party shall be entitled to seek injunctive relief in addition to any other remedies available at law or equity.")

    # General
    add_heading(doc, "9. GENERAL PROVISIONS", 1)
    add_paragraph(doc, "9.1 Governing Law. This Agreement shall be governed by the laws of England and Wales.")
    add_paragraph(doc, "9.2 Entire Agreement. This Agreement constitutes the entire agreement between the Parties regarding the subject matter hereof.")
    add_paragraph(doc, "9.3 Amendment. This Agreement may only be modified by a written agreement signed by both Parties.")

    # Signatures
    doc.add_paragraph()
    add_paragraph(doc, "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "ACME CORPORATION"
    table.cell(0, 1).text = "GLOBALPARTNERS LTD."
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name: John Smith"
    table.cell(2, 1).text = "Name: William Brown"
    table.cell(3, 0).text = "Title: CEO"
    table.cell(3, 1).text = "Title: Managing Director"

    filepath = os.path.join(OUTPUT_DIR, "Acme_GlobalPartners_NDA_2025.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_software_license():
    """Create a Software License Agreement with SLAs."""
    doc = Document()

    # Title
    title = doc.add_heading("SOFTWARE LICENSE AND SUPPORT AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_paragraph(doc, "This Software License and Support Agreement ('Agreement') is made effective as of March 1, 2025 ('Effective Date') between:")
    doc.add_paragraph()
    add_paragraph(doc, "CloudSoft Inc., a Delaware corporation ('Licensor')")
    add_paragraph(doc, "AND")
    add_paragraph(doc, "Acme Corporation, a Delaware corporation ('Licensee')")

    # Grant of License
    add_heading(doc, "1. GRANT OF LICENSE", 1)
    add_paragraph(doc, "1.1 License Grant. Subject to the terms of this Agreement, Licensor grants Licensee a non-exclusive, non-transferable license to use the CloudSoft Enterprise Platform ('Software') for Licensee's internal business operations.")
    add_paragraph(doc, "1.2 License Scope. The license permits use by up to 500 named users and unlimited API calls per month.")
    add_paragraph(doc, "1.3 Restrictions. Licensee shall not: (a) sublicense, sell, or transfer the Software; (b) reverse engineer or decompile the Software; (c) remove any proprietary notices; (d) use the Software for any unlawful purpose.")

    # Fees
    add_heading(doc, "2. FEES AND PAYMENT", 1)
    add_paragraph(doc, "2.1 License Fee. Licensee shall pay an annual license fee of $250,000 USD, payable in advance on each anniversary of the Effective Date.")
    add_paragraph(doc, "2.2 Support Fee. Annual support and maintenance fee of $50,000 USD is included in the license fee for the first year. Subsequent years will be billed at 20% of the then-current license fee.")
    add_paragraph(doc, "2.3 Late Payment. Any amounts not paid when due shall bear interest at 1.5% per month or the maximum rate permitted by law, whichever is less.")

    # Service Level Agreement
    add_heading(doc, "3. SERVICE LEVEL AGREEMENT", 1)
    add_paragraph(doc, "3.1 Availability. Licensor guarantees 99.9% uptime availability for the hosted Software, measured monthly, excluding scheduled maintenance windows.")
    add_paragraph(doc, "3.2 Response Times. Licensor shall respond to support requests as follows:")
    add_paragraph(doc, "   - Critical (System Down): Response within 1 hour, Resolution within 4 hours")
    add_paragraph(doc, "   - High (Major Feature Impaired): Response within 4 hours, Resolution within 8 hours")
    add_paragraph(doc, "   - Medium (Minor Feature Impaired): Response within 8 hours, Resolution within 24 hours")
    add_paragraph(doc, "   - Low (General Questions): Response within 24 hours, Resolution within 72 hours")
    add_paragraph(doc, "3.3 Service Credits. If Licensor fails to meet the availability guarantee, Licensee shall receive service credits as follows:")
    add_paragraph(doc, "   - 99.0% - 99.9%: 10% credit of monthly fee")
    add_paragraph(doc, "   - 95.0% - 99.0%: 25% credit of monthly fee")
    add_paragraph(doc, "   - Below 95.0%: 50% credit of monthly fee")
    add_paragraph(doc, "3.4 Credit Cap. Total service credits shall not exceed 50% of monthly fees in any given month.")

    # Support Services
    add_heading(doc, "4. SUPPORT SERVICES", 1)
    add_paragraph(doc, "4.1 Scope. Support services include: (a) telephone and email technical support during business hours (8am-8pm EST, Monday-Friday); (b) access to online knowledge base and documentation; (c) bug fixes and security patches; (d) minor version updates.")
    add_paragraph(doc, "4.2 Exclusions. Support does not include: (a) customization services; (b) training; (c) issues caused by Licensee modifications; (d) third-party software integration issues.")

    # Data Security
    add_heading(doc, "5. DATA SECURITY", 1)
    add_paragraph(doc, "5.1 Security Measures. Licensor shall implement and maintain appropriate security measures including: AES-256 encryption at rest, TLS 1.3 encryption in transit, multi-factor authentication, and annual SOC 2 Type II audits.")
    add_paragraph(doc, "5.2 Data Backup. Licensor shall perform daily backups with 30-day retention and maintain disaster recovery capabilities with RPO of 4 hours and RTO of 8 hours.")
    add_paragraph(doc, "5.3 Breach Notification. Licensor shall notify Licensee of any security breach affecting Licensee data within 48 hours of discovery.")

    # Term and Termination
    add_heading(doc, "6. TERM AND TERMINATION", 1)
    add_paragraph(doc, "6.1 Term. This Agreement shall commence on the Effective Date and continue for an initial term of three (3) years. The Agreement shall automatically renew for successive one (1) year terms unless either party provides written notice of non-renewal at least sixty (60) days prior to the end of the then-current term.")
    add_paragraph(doc, "6.2 Termination for Breach. Either party may terminate this Agreement upon thirty (30) days written notice if the other party materially breaches this Agreement and fails to cure such breach within the notice period.")
    add_paragraph(doc, "6.3 Effect of Termination. Upon termination: (a) all licenses granted herein shall immediately terminate; (b) Licensee shall cease use of the Software; (c) Licensor shall provide Licensee's data in a standard format within 30 days.")

    # Limitation of Liability
    add_heading(doc, "7. LIMITATION OF LIABILITY", 1)
    add_paragraph(doc, "7.1 EXCLUSION OF DAMAGES. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOST PROFITS OR LOST DATA.")
    add_paragraph(doc, "7.2 CAP ON LIABILITY. EACH PARTY'S TOTAL CUMULATIVE LIABILITY SHALL NOT EXCEED THE FEES PAID OR PAYABLE BY LICENSEE IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM.")
    add_paragraph(doc, "7.3 EXCEPTIONS. The limitations in this Section 7 shall not apply to: (a) breaches of confidentiality; (b) infringement of intellectual property rights; (c) gross negligence or willful misconduct.")

    # General
    add_heading(doc, "8. GENERAL PROVISIONS", 1)
    add_paragraph(doc, "8.1 Governing Law. This Agreement shall be governed by the laws of the State of Delaware.")
    add_paragraph(doc, "8.2 Dispute Resolution. Any disputes shall first be subject to good faith negotiation. If unresolved after 30 days, disputes shall be submitted to binding arbitration under AAA Commercial Rules.")
    add_paragraph(doc, "8.3 Force Majeure. Neither party shall be liable for delays caused by events beyond its reasonable control.")
    add_paragraph(doc, "8.4 Entire Agreement. This Agreement constitutes the entire agreement between the parties.")

    # Signatures
    doc.add_paragraph()
    add_paragraph(doc, "IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "CLOUDSOFT INC."
    table.cell(0, 1).text = "ACME CORPORATION"
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name: Sarah Johnson"
    table.cell(2, 1).text = "Name: John Smith"
    table.cell(3, 0).text = "Title: VP of Sales"
    table.cell(3, 1).text = "Title: CTO"

    filepath = os.path.join(OUTPUT_DIR, "CloudSoft_Acme_License_2025.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_consulting_agreement():
    """Create a Consulting Agreement for TechStart tenant."""
    doc = Document()

    # Title
    title = doc.add_heading("CONSULTING SERVICES AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_paragraph(doc, "This Consulting Services Agreement ('Agreement') is entered into as of January 20, 2025 ('Effective Date') by and between:")
    doc.add_paragraph()
    add_paragraph(doc, "TechStart Inc., a California corporation ('Client')")
    add_paragraph(doc, "AND")
    add_paragraph(doc, "Strategic Advisors LLC, a Delaware limited liability company ('Consultant')")

    # Services
    add_heading(doc, "1. CONSULTING SERVICES", 1)
    add_paragraph(doc, "1.1 Scope. Consultant shall provide strategic consulting services including: (a) market analysis and competitive intelligence; (b) go-to-market strategy development; (c) product positioning and messaging; (d) investor presentation preparation.")
    add_paragraph(doc, "1.2 Deliverables. Consultant shall deliver: (a) Market Analysis Report within 30 days; (b) Strategic Plan within 60 days; (c) Investor Deck within 45 days.")
    add_paragraph(doc, "1.3 Time Commitment. Consultant shall dedicate approximately 40 hours per month to this engagement.")

    # Compensation
    add_heading(doc, "2. COMPENSATION", 1)
    add_paragraph(doc, "2.1 Fees. Client shall pay Consultant a monthly retainer of $15,000 USD, plus a success fee of 1% of any funding raised during the term.")
    add_paragraph(doc, "2.2 Expenses. Client shall reimburse Consultant for pre-approved travel and out-of-pocket expenses.")
    add_paragraph(doc, "2.3 Invoicing. Consultant shall invoice monthly, with payment due within 15 days of invoice date.")

    # Term
    add_heading(doc, "3. TERM AND TERMINATION", 1)
    add_paragraph(doc, "3.1 Term. This Agreement shall be effective for six (6) months from the Effective Date.")
    add_paragraph(doc, "3.2 Termination. Either party may terminate with 14 days written notice. Upon termination, Client shall pay for all services rendered through the termination date.")

    # Confidentiality
    add_heading(doc, "4. CONFIDENTIALITY", 1)
    add_paragraph(doc, "4.1 Obligations. Consultant shall maintain strict confidentiality of all Client information and shall not disclose to any third party without prior written consent.")
    add_paragraph(doc, "4.2 Non-Compete. During the term and for 12 months thereafter, Consultant shall not provide similar services to direct competitors of Client.")

    # Intellectual Property
    add_heading(doc, "5. INTELLECTUAL PROPERTY", 1)
    add_paragraph(doc, "All deliverables and work product created under this Agreement shall be the exclusive property of Client upon payment.")

    # Liability
    add_heading(doc, "6. LIABILITY", 1)
    add_paragraph(doc, "Consultant's total liability shall not exceed the fees paid under this Agreement. Neither party shall be liable for indirect or consequential damages.")

    # General
    add_heading(doc, "7. GENERAL", 1)
    add_paragraph(doc, "7.1 Independent Contractor. Consultant is an independent contractor, not an employee of Client.")
    add_paragraph(doc, "7.2 Governing Law. This Agreement shall be governed by California law.")

    # Signatures
    doc.add_paragraph()
    add_paragraph(doc, "IN WITNESS WHEREOF, the parties have executed this Agreement.")
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "TECHSTART INC."
    table.cell(0, 1).text = "STRATEGIC ADVISORS LLC"
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name: Michael Chen"
    table.cell(2, 1).text = "Name: David Wilson"
    table.cell(3, 0).text = "Title: CEO"
    table.cell(3, 1).text = "Title: Managing Partner"

    filepath = os.path.join(OUTPUT_DIR, "TechStart_StrategicAdvisors_Consulting_2025.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


def create_employment_agreement():
    """Create an Employment Agreement for TechStart tenant."""
    doc = Document()

    # Title
    title = doc.add_heading("EMPLOYMENT AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_paragraph(doc, "This Employment Agreement ('Agreement') is entered into as of February 15, 2025 ('Effective Date') by and between:")
    doc.add_paragraph()
    add_paragraph(doc, "TechStart Inc., a California corporation ('Company')")
    add_paragraph(doc, "AND")
    add_paragraph(doc, "Emily Rodriguez ('Employee')")

    # Position
    add_heading(doc, "1. POSITION AND DUTIES", 1)
    add_paragraph(doc, "1.1 Position. Company hereby employs Employee as Vice President of Engineering, reporting to the Chief Executive Officer.")
    add_paragraph(doc, "1.2 Duties. Employee shall perform all duties customarily associated with such position and such other duties as may be assigned.")
    add_paragraph(doc, "1.3 Best Efforts. Employee agrees to devote their full business time and best efforts to the Company.")

    # Compensation
    add_heading(doc, "2. COMPENSATION", 1)
    add_paragraph(doc, "2.1 Base Salary. Employee shall receive an annual base salary of $275,000 USD, payable bi-weekly.")
    add_paragraph(doc, "2.2 Bonus. Employee shall be eligible for an annual performance bonus of up to 30% of base salary, based on achievement of performance targets.")
    add_paragraph(doc, "2.3 Equity. Employee shall be granted stock options to purchase 50,000 shares of Company common stock, vesting over 4 years with a 1-year cliff.")
    add_paragraph(doc, "2.4 Benefits. Employee shall be eligible for Company's standard benefits including health insurance, 401(k), and 20 days PTO.")

    # Term
    add_heading(doc, "3. TERM AND TERMINATION", 1)
    add_paragraph(doc, "3.1 At-Will Employment. Employment is at-will and may be terminated by either party at any time.")
    add_paragraph(doc, "3.2 Severance. If Company terminates Employee without Cause, Employee shall receive 6 months base salary as severance, subject to signing a release.")
    add_paragraph(doc, "3.3 Change of Control. If Employee is terminated within 12 months of a Change of Control, Employee shall receive 12 months severance plus accelerated vesting.")

    # Confidentiality
    add_heading(doc, "4. CONFIDENTIALITY AND IP", 1)
    add_paragraph(doc, "4.1 Confidentiality. Employee shall maintain confidentiality of all Company proprietary information during and after employment.")
    add_paragraph(doc, "4.2 Inventions. All inventions, developments, and work product created during employment shall be Company property.")
    add_paragraph(doc, "4.3 Non-Solicitation. For 12 months after termination, Employee shall not solicit Company employees or customers.")

    # General
    add_heading(doc, "5. GENERAL PROVISIONS", 1)
    add_paragraph(doc, "5.1 Governing Law. This Agreement shall be governed by California law.")
    add_paragraph(doc, "5.2 Arbitration. Any disputes shall be resolved by binding arbitration in San Francisco.")
    add_paragraph(doc, "5.3 Entire Agreement. This Agreement constitutes the entire agreement between the parties.")

    # Signatures
    doc.add_paragraph()
    add_paragraph(doc, "IN WITNESS WHEREOF, the parties have executed this Agreement.")
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "TECHSTART INC."
    table.cell(0, 1).text = "EMPLOYEE"
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name: Michael Chen"
    table.cell(2, 1).text = "Name: Emily Rodriguez"
    table.cell(3, 0).text = "Title: CEO"
    table.cell(3, 1).text = ""

    filepath = os.path.join(OUTPUT_DIR, "TechStart_Rodriguez_Employment_2025.docx")
    doc.save(filepath)
    print(f"Created: {filepath}")
    return filepath


if __name__ == "__main__":
    print("Generating test contract documents...")
    print()

    # Create contracts for Acme Corp (admin user)
    print("=== For Acme Corp (admin user) ===")
    create_msa()
    create_nda()
    create_software_license()

    print()

    # Create contracts for TechStart (techstart_admin user)
    print("=== For TechStart (techstart_admin user) ===")
    create_consulting_agreement()
    create_employment_agreement()

    print()
    print("Done! Upload these contracts to test the platform.")
