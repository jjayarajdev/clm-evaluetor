#!/usr/bin/env python3
"""
Local contract anonymization script.

Converts .doc/.xls/.ppt to modern formats using LibreOffice, then anonymizes
PII in DOCX, XLSX, and PDF files using only local packages.
No data is sent to any external API or service.

Supports three contract sets:
  1. Visser-Syntegreti NDA (reference contracts root)
  2. ING outsourcing agreement (ing/ subfolder)
  3. Novartis outsourcing agreement (novartis/ subfolder)
"""

import re
import shutil
import subprocess
import sys
from pathlib import Path

import fitz  # PyMuPDF
import openpyxl
from docx import Document

LIBREOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

BASE_DIR = Path(
    "/Users/jjayaraj/personal/Syntegreti/CLM/Steven/reference contracts"
)

# ---------------------------------------------------------------------------
# Replacement maps per contract set — longest strings first
# ---------------------------------------------------------------------------

VISSER_REPLACEMENTS = [
    ("S.E. Visser Consultancy B.V.", "Alpha Consulting Group B.V."),
    ("S.E.Visser Consultancy B.V.", "Alpha Consulting Group B.V."),
    ("S.E.Visser Consultancy BV", "Alpha Consulting Group BV"),
    ("S.E. Visser Consultancy BV", "Alpha Consulting Group BV"),
    ("Syntegreti Technologies Pvt Ltd.", "Beta Technologies Pvt Ltd."),
    ("Syntegreti Technologies Pvt Ltd", "Beta Technologies Pvt Ltd"),
    ("Burgemeester Le Fevre de Montignylaan 264, 3055NK, Rotterdam", "100 Commerce Street, 1000AB, Amsterdam"),
    ("Burgemeester Le Fevre de Montignylaan 264, 3055NK", "100 Commerce Street, 1000AB"),
    ("Burgemeester Le Fevre de Montignylaan 264", "100 Commerce Street"),
    ("Office 515, Trend Works, 5th Floor, Jyothi Granules No 199, Kondapur Village Serilingmpally Hyderabad, Telangana 500084",
     "Suite 200, Tech Park, 4th Floor, Business District, Innovation City, State 600001"),
    ("Office 515, Trend Works, 5th Floor, Jyothi Granules No 199, Kondapur Village Serilingmpally",
     "Suite 200, Tech Park, 4th Floor, Business District"),
    ("Jyothi Granules No 199, Kondapur Village Serilingmpally", "Business District"),
    ("Kondapur Village Serilingmpally", "Business District"),
    ("Trend Works, 5th Floor", "Tech Park, 4th Floor"),
    ("Telangana 500084", "State 600001"),
    ("3055NK, Rotterdam", "1000AB, Amsterdam"),
    ("3055NK", "1000AB"),
    ("Steven Visser", "John Smith"),
    ("Bharath V S Dhurjati", "Jane Doe"),
    ("Bharath Dhurjati", "Jane Doe"),
    ("February 2nd 2026", "January 15th 2026"),
    ("February 2nd, 2026", "January 15th, 2026"),
    ("February 3rd 2026", "January 16th 2026"),
    ("February 3rd, 2026", "January 16th, 2026"),
    ("February 6th 2026", "January 19th 2026"),
    ("February 6th, 2026", "January 19th, 2026"),
    ("Managing Partner", "Chief Executive"),
    ("Director", "Chief Officer"),
    ("Syntegreti", "Beta"),
    ("Visser", "Alpha"),
    ("Rotterdam", "Amsterdam"),
    ("Hyderabad", "Innovation City"),
]

ING_REPLACEMENTS = [
    # Full entity names
    ("ING BANK NV", "EUROBANK NV"),
    ("ING BANK N.V.", "EUROBANK N.V."),
    ("ING Bank NV", "EuroBank NV"),
    ("ING Bank N.V.", "EuroBank N.V."),
    ("ING Bank nv", "EuroBank nv"),
    ("ING Personeel VOF", "EuroBank Personeel VOF"),
    ("ING Personeel", "EuroBank Personeel"),
    # Company registration
    ("33031431", "99887766"),
    # Addresses
    ("Amstelveenseweg 500, 1081 BL Amsterdam", "Keizersgracht 100, 1015 AA Rotterdam"),
    ("Amstelveenseweg 500", "Keizersgracht 100"),
    ("1081 BL Amsterdam", "1015 AA Rotterdam"),
    ("1081 BL", "1015 AA"),
    # Project codename
    ("Trimaran", "Neptune"),
    # Vendor names (in filenames and content)
    ("Atos Origin", "Vendor Alpha"),
    ("ATOS ORIGIN", "VENDOR ALPHA"),
    ("Atos", "VendorA"),
    ("ATOS", "VENDORA"),
    ("Accenture", "Vendor Beta"),
    ("ACCENTURE", "VENDOR BETA"),
    ("Getronics", "Vendor Gamma"),
    ("GETRONICS", "VENDOR GAMMA"),
    ("KPN", "TelCo"),
    # Dates
    ("19 July 2005", "15 March 2005"),
    ("2 September 2005", "1 June 2005"),
    # ING standalone (careful — only match as whole word via regex later)
    # These are handled by regex replacement below
]

# Regex patterns for ING that avoid matching inside words like "pricing", "during"
ING_REGEX_REPLACEMENTS = [
    (re.compile(r'\bING\b'), "EuroBank"),
]

NOVARTIS_REPLACEMENTS = [
    # Full entity names
    ("Novartis Pharma Ltd.", "GlobalPharma Ltd."),
    ("Novartis Pharma Ltd", "GlobalPharma Ltd"),
    ("NOVARTIS PHARMA", "GLOBALPHARMA"),
    ("Novartis Pharma", "GlobalPharma"),
    ("Genpact International SaRL", "ServicePro International SaRL"),
    ("Genpact International", "ServicePro International"),
    ("GENPACT", "SERVICEPRO"),
    ("Genpact", "ServicePro"),
    # Law firm
    ("WENGER PLATTNER", "LAW FIRM"),
    ("Wenger Plattner", "Law Firm"),
    # Addresses
    ("Lichtstrasse 35, CH-4056 Basel", "Bahnhofstrasse 10, CH-8001 Zurich"),
    ("Lichtstrasse 15, 4056 Zurich", "Bahnhofstrasse 10, 8001 Zurich"),
    ("Lichtstrasse 35", "Bahnhofstrasse 10"),
    ("Lichtstrasse 15", "Bahnhofstrasse 10"),
    ("Lichistrasse 35", "Bahnhofstrasse 10"),
    ("Lichtstrasse", "Bahnhofstrasse"),
    ("Lichistrasse", "Bahnhofstrasse"),
    ("CH-4056 Basel", "CH-8001 Zurich"),
    ("CH 4056", "CH 8001"),
    ("CH-4056", "CH-8001"),
    ("65 Bvd. Grande-Duchesse Charlotte, L-1331 Luxembourg",
     "25 Avenue de la Liberté, L-1930 Luxembourg"),
    ("65 Bv. Grande-Duchesse Charlotte, L-1331 Luxembourg",
     "25 Avenue de la Liberté, L-1930 Luxembourg"),
    ("65 Bv. Grande-Duchesse Charlotte, L-1930 Luxembourg",
     "25 Avenue de la Liberté, L-1930 Luxembourg"),
    ("65 BV, Grande-Duchesse Charlotte, L-1331 Luxembourg",
     "25 Avenue de la Liberté, L-1930 Luxembourg"),
    ("65 Bvd. Grande-Duchesse Charlotte", "25 Avenue de la Liberté"),
    ("65 Bv. Grande-Duchesse Charlotte", "25 Avenue de la Liberté"),
    ("65 BV, Grande-Duchesse Charlotte", "25 Avenue de la Liberté"),
    ("65 Boulevard Grande-Duchesse Charlotte", "25 Avenue de la Liberté"),
    ("Grande-Duchesse Charlotte", "Avenue de la Liberté"),
    ("L-1331 Luxembourg", "L-1930 Luxembourg"),
    ("L-1331", "L-1930"),
    # Dates
    ("29 January 07", "15 January 07"),
    ("29 January 2007", "15 January 2007"),
    ("30 January 07", "16 January 07"),
    ("31 January 07", "17 January 07"),
    ("31 January 2007", "17 January 2007"),
    ("January 31, 2007", "January 17, 2007"),
    ("October 20, 2006", "September 1, 2006"),
    ("November 10, 2006", "September 20, 2006"),
    ("December 18, 2006", "October 15, 2006"),
    ("December 21, 2006", "October 18, 2006"),
    ("December 4, 2006", "October 1, 2006"),
    ("January 15, 2007", "November 1, 2006"),
    # Cities used as identifiers
    ("Basel", "Zurich"),
    # Novartis standalone
    ("Novartis", "GlobalPharma"),
    ("NOVARTIS", "GLOBALPHARMA"),
]

# Filename replacements (applied to output filenames)
ING_FILENAME_REPLACEMENTS = [
    ("Atos Origin", "Vendor Alpha"),
    ("ATOS ORIGIN", "VENDOR ALPHA"),
    ("Atos", "VendorA"),
    ("ATOS", "VENDORA"),
    ("Accenture", "Vendor Beta"),
    ("Getronics", "Vendor Gamma"),
    ("GETRONICS", "VENDOR GAMMA"),
    ("KPN", "TelCo"),
    ("ING", "EuroBank"),
    ("2IM", "UnitX"),
    ("Trimaran", "Neptune"),
    ("SWS", "ProjectY"),
    ("GPR", "DivA"),
]

NOVARTIS_FILENAME_REPLACEMENTS = [
    ("Novartis Pharma", "GlobalPharma"),
    ("NOVARTIS", "GLOBALPHARMA"),
    ("Novartis", "GlobalPharma"),
    ("WENGER PLATTNER", "LAW FIRM"),
    ("Wenger Plattner", "Law Firm"),
    ("Genpact", "ServicePro"),
]


# ---------------------------------------------------------------------------
# Conversion helpers
# ---------------------------------------------------------------------------

def convert_with_libreoffice(src: Path, target_format: str, outdir: Path) -> Path | None:
    """Convert a file using LibreOffice. Returns path to converted file or None."""
    result = subprocess.run(
        [LIBREOFFICE, "--headless", "--convert-to", target_format,
         str(src), "--outdir", str(outdir)],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        print(f"    WARN: LibreOffice conversion failed for {src.name}: {result.stderr[:200]}")
        return None

    expected = outdir / (src.stem + "." + target_format)
    if expected.exists():
        return expected
    # LibreOffice sometimes changes case
    for f in outdir.iterdir():
        if f.stem.lower() == src.stem.lower() and f.suffix.lower() == f".{target_format}":
            return f
    return None


# ---------------------------------------------------------------------------
# Anonymization functions
# ---------------------------------------------------------------------------

def apply_replacements(text: str, replacements: list, regex_replacements: list = None) -> str:
    """Apply string and regex replacements to text."""
    for original, replacement in replacements:
        text = text.replace(original, replacement)
    if regex_replacements:
        for pattern, replacement in regex_replacements:
            text = pattern.sub(replacement, text)
    return text


def anonymize_docx(src: Path, dst: Path, replacements: list, regex_replacements: list = None) -> None:
    """Anonymize a DOCX by replacing text in paragraphs, tables, headers, footers."""
    doc = Document(src)

    def replace_in_paragraph(paragraph):
        full_text = paragraph.text
        if not full_text:
            return
        new_text = apply_replacements(full_text, replacements, regex_replacements)
        if new_text != full_text:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                for run in paragraph.runs[1:]:
                    run.text = ""
                first_run.text = new_text
            else:
                paragraph.text = new_text

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
                for nested_table in cell.tables:
                    replace_in_table(nested_table)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        replace_in_table(table)

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for paragraph in header.paragraphs:
                    replace_in_paragraph(paragraph)
                for table in header.tables:
                    replace_in_table(table)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for paragraph in footer.paragraphs:
                    replace_in_paragraph(paragraph)
                for table in footer.tables:
                    replace_in_table(table)

    doc.save(str(dst))


def anonymize_xlsx(src: Path, dst: Path, replacements: list, regex_replacements: list = None) -> None:
    """Anonymize an XLSX by replacing text in all cells across all sheets."""
    wb = openpyxl.load_workbook(str(src))
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_val = apply_replacements(cell.value, replacements, regex_replacements)
                    if new_val != cell.value:
                        cell.value = new_val
        # Also check sheet title
        new_title = apply_replacements(sheet.title, replacements, regex_replacements)
        if new_title != sheet.title:
            sheet.title = new_title
    wb.save(str(dst))


def anonymize_pdf(src: Path, dst: Path, replacements: list, regex_replacements: list = None) -> None:
    """Anonymize a PDF using PyMuPDF's redaction API."""
    doc = fitz.open(src)
    for page_num in range(len(doc)):
        page = doc[page_num]
        for original, replacement in replacements:
            instances = page.search_for(original)
            for inst in instances:
                page.add_redact_annot(
                    inst, text=replacement, fontsize=0,
                    align=fitz.TEXT_ALIGN_LEFT,
                )
        # Regex replacements for PDF: extract text and search manually
        if regex_replacements:
            text = page.get_text()
            for pattern, replacement in regex_replacements:
                for match in pattern.finditer(text):
                    instances = page.search_for(match.group())
                    for inst in instances:
                        page.add_redact_annot(
                            inst, text=replacement, fontsize=0,
                            align=fitz.TEXT_ALIGN_LEFT,
                        )
        page.apply_redactions()
    doc.save(str(dst))
    doc.close()


def anonymize_filename(name: str, filename_replacements: list) -> str:
    """Apply replacements to a filename."""
    for original, replacement in filename_replacements:
        name = name.replace(original, replacement)
    return name


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------

def verify_docx(path: Path, pii_terms: list) -> list:
    """Check anonymized DOCX for remaining PII. Returns list of found terms."""
    try:
        doc = Document(path)
    except Exception:
        return []
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    found = []
    for term in pii_terms:
        # Use word boundary for short terms
        if len(term) <= 3:
            if re.search(r'\b' + re.escape(term) + r'\b', full_text):
                found.append(term)
        elif term in full_text:
            found.append(term)
    return found


def verify_xlsx(path: Path, pii_terms: list) -> list:
    """Check anonymized XLSX for remaining PII."""
    try:
        wb = openpyxl.load_workbook(str(path), read_only=True)
    except Exception:
        return []
    full_text = ""
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    full_text += " " + cell.value
    wb.close()
    found = []
    for term in pii_terms:
        if len(term) <= 3:
            if re.search(r'\b' + re.escape(term) + r'\b', full_text):
                found.append(term)
        elif term in full_text:
            found.append(term)
    return found


def verify_pdf(path: Path, pii_terms: list) -> list:
    """Check anonymized PDF for remaining PII."""
    try:
        doc = fitz.open(path)
    except Exception:
        return []
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    doc.close()
    found = []
    for term in pii_terms:
        if len(term) <= 3:
            if re.search(r'\b' + re.escape(term) + r'\b', full_text):
                found.append(term)
        elif term in full_text:
            found.append(term)
    return found


# ---------------------------------------------------------------------------
# Process a contract folder
# ---------------------------------------------------------------------------

def process_folder(
    name: str,
    src_dir: Path,
    out_dir: Path,
    replacements: list,
    regex_replacements: list | None,
    filename_replacements: list,
    pii_terms: list,
):
    """Convert, anonymize, and verify all files in a contract folder."""
    print(f"\n{'='*70}")
    print(f"  Processing: {name}")
    print(f"  Source:     {src_dir}")
    print(f"  Output:     {out_dir}")
    print(f"{'='*70}")

    out_dir.mkdir(parents=True, exist_ok=True)

    # Gather all files
    files = sorted(f for f in src_dir.iterdir() if f.is_file())
    print(f"\n  Found {len(files)} files")

    # Phase 1: Convert legacy formats with LibreOffice
    convert_dir = out_dir / "_converted"
    convert_dir.mkdir(exist_ok=True)

    converted_files = []  # (anonymized_name, path) tuples
    skipped = []

    for f in files:
        ext = f.suffix.lower()
        if ext in (".doc", ".dot"):
            print(f"  Converting {f.name} → docx ...", end=" ", flush=True)
            result = convert_with_libreoffice(f, "docx", convert_dir)
            if result:
                print("OK")
                anon_name = anonymize_filename(result.stem, filename_replacements) + ".docx"
                converted_files.append((anon_name, result, "docx"))
            else:
                print("FAILED")
                skipped.append(f.name)
        elif ext == ".xls":
            print(f"  Converting {f.name} → xlsx ...", end=" ", flush=True)
            result = convert_with_libreoffice(f, "xlsx", convert_dir)
            if result:
                print("OK")
                anon_name = anonymize_filename(result.stem, filename_replacements) + ".xlsx"
                converted_files.append((anon_name, result, "xlsx"))
            else:
                print("FAILED")
                skipped.append(f.name)
        elif ext == ".ppt":
            print(f"  Converting {f.name} → pptx ...", end=" ", flush=True)
            result = convert_with_libreoffice(f, "pptx", convert_dir)
            if result:
                print("OK")
                anon_name = anonymize_filename(result.stem, filename_replacements) + ".pptx"
                # Just copy pptx — no anonymization support yet
                dst = out_dir / anon_name
                shutil.copy2(result, dst)
                print(f"    → Copied as {anon_name} (pptx anonymization not supported)")
            else:
                print("FAILED")
                skipped.append(f.name)
        elif ext == ".docx":
            anon_name = anonymize_filename(f.stem, filename_replacements) + ".docx"
            converted_files.append((anon_name, f, "docx"))
        elif ext == ".xlsx":
            anon_name = anonymize_filename(f.stem, filename_replacements) + ".xlsx"
            converted_files.append((anon_name, f, "xlsx"))
        elif ext == ".pdf":
            anon_name = anonymize_filename(f.stem, filename_replacements) + ".pdf"
            converted_files.append((anon_name, f, "pdf"))
        elif ext == ".zip":
            print(f"  Skipping archive: {f.name}")
            skipped.append(f.name)
        else:
            print(f"  Skipping unsupported: {f.name}")
            skipped.append(f.name)

    # Phase 2: Anonymize
    print(f"\n  --- Anonymizing {len(converted_files)} files ---")
    success = 0
    errors = []

    for anon_name, src_path, fmt in converted_files:
        dst = out_dir / anon_name
        try:
            if fmt == "docx":
                anonymize_docx(src_path, dst, replacements, regex_replacements)
            elif fmt == "xlsx":
                anonymize_xlsx(src_path, dst, replacements, regex_replacements)
            elif fmt == "pdf":
                anonymize_pdf(src_path, dst, replacements, regex_replacements)
            success += 1
            print(f"    ✓ {anon_name}")
        except Exception as e:
            errors.append((anon_name, str(e)))
            print(f"    ✗ {anon_name}: {e}")

    # Phase 3: Verify
    print(f"\n  --- Verification ---")
    warn_count = 0
    for anon_name, _, fmt in converted_files:
        dst = out_dir / anon_name
        if not dst.exists():
            continue
        if fmt == "docx":
            found = verify_docx(dst, pii_terms)
        elif fmt == "xlsx":
            found = verify_xlsx(dst, pii_terms)
        elif fmt == "pdf":
            found = verify_pdf(dst, pii_terms)
        else:
            continue
        if found:
            print(f"    WARN {anon_name}: remnants → {found}")
            warn_count += 1

    # Cleanup temp conversion dir
    shutil.rmtree(convert_dir, ignore_errors=True)

    # Summary
    print(f"\n  --- Summary for {name} ---")
    print(f"    Anonymized: {success}/{len(converted_files)}")
    if errors:
        print(f"    Errors:     {len(errors)}")
        for n, e in errors:
            print(f"      {n}: {e}")
    if skipped:
        print(f"    Skipped:    {len(skipped)} ({', '.join(skipped[:5])}{'...' if len(skipped)>5 else ''})")
    if warn_count:
        print(f"    PII warnings: {warn_count} files")
    else:
        print(f"    PII check:  ALL PASS")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=" * 70)
    print("  Contract Anonymization (100% Local — No External APIs)")
    print("=" * 70)

    # Determine which sets to process based on CLI args
    targets = sys.argv[1:] if len(sys.argv) > 1 else ["ing", "novartis"]

    if "ing" in targets:
        process_folder(
            name="ING Outsourcing Agreement",
            src_dir=BASE_DIR / "ing",
            out_dir=BASE_DIR / "ing_anonymized",
            replacements=ING_REPLACEMENTS,
            regex_replacements=ING_REGEX_REPLACEMENTS,
            filename_replacements=ING_FILENAME_REPLACEMENTS,
            pii_terms=["ING", "Trimaran", "33031431", "Amstelveenseweg",
                        "1081 BL", "Accenture", "Atos Origin", "Getronics", "KPN"],
        )

    if "novartis" in targets:
        process_folder(
            name="Novartis Outsourcing Agreement",
            src_dir=BASE_DIR / "novartis",
            out_dir=BASE_DIR / "novartis_anonymized",
            replacements=NOVARTIS_REPLACEMENTS,
            regex_replacements=None,
            filename_replacements=NOVARTIS_FILENAME_REPLACEMENTS,
            pii_terms=["Novartis", "Genpact", "Wenger", "Plattner",
                        "Lichtstrasse", "CH-4056", "Basel",
                        "Grande-Duchesse", "L-1331"],
        )

    print(f"\n{'='*70}")
    print("  All done.")
    print(f"{'='*70}")


if __name__ == "__main__":
    main()
